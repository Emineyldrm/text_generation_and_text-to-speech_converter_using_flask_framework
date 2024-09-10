from typing import Any
from flask import Flask,render_template, request, redirect, url_for, render_template_string, send_file
from text_generation import generate_text_from_question
from text_to_speech import create_audio #, ses_dosyasi
import os
from pdf_okuma import pdf_oku_ve_metni_al
from silme import eski_dosyalari_temizle
from docx import Document
from io import BytesIO
#from gtts.lang import tts_langs


app=Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

@app.route('/')
def home():#ana_sayfa
    return render_template('index.html')
 

@app.route('/pdf')
def pdf():
    return render_template('pdf.html') 

@app.route('/anasayfa')
def gelistirici():
    return render_template('anasayfa.html') 


@app.route('/submit', methods=['POST'])
def submit(): 
    language=request.form['language']
    soru = request.form['soru']
    transcript = request.form.get('transcript')    #yeni eklendi
    #response = generate_text_from_question(soru)
    #response = soru                                                    #BURASI DEĞİŞECEK
   
    if soru:
        #response = soru
        response = generate_text_from_question(soru)
    elif transcript:
        soru = transcript
        response = generate_text_from_question(soru)        
        #response = transcript
    else:
        response = "No input received."
    #create_audio(response,language)
    eski_dosyalari_temizle()
    unique_filename = create_audio(response, language)        
    return render_template('index.html', soru=soru, response=response,audio_file=unique_filename)


@app.route('/upload', methods=['POST'])
def upload_pdf():
    if 'pdf_dosyasi' not in request.files:
        return redirect(url_for('pdf'))
    
    dosya = request.files['pdf_dosyasi']
    

    if dosya.filename == '':
        return redirect(url_for('pdf'))
    
    try:
        baslangic_sayfa = int(request.form.get('baslangic_sayfa', 1)) - 1  # Sıfırdan başlatmak için 1 çıkar
        bitis_sayfa = int(request.form.get('bitis_sayfa'))  # Sayfa numaraları 1'den başlıyor
        language = request.form.get('language')
    except ValueError:
        return redirect(url_for('pdf'))

    if dosya and dosya.filename.endswith('.pdf'):
        dosya_yolu = os.path.join(app.config['UPLOAD_FOLDER'], dosya.filename)
        dosya.save(dosya_yolu)
        pdf_icerik = pdf_oku_ve_metni_al(dosya_yolu, baslangic_sayfa, bitis_sayfa)
        os.remove(dosya_yolu)  # Dosyayı işlem sonrası sil
        
        # Generate audio from the extracted text
        eski_dosyalari_temizle()
        unique_filename = create_audio(pdf_icerik, language)
        return render_template('pdf.html', pdf_icerik=pdf_icerik,audio_file=unique_filename)
    
    return redirect(url_for('pdf'))

@app.route('/download_docx')
def download_docx():
    pdf_content = request.args.get('pdf_content', '')
    
    # Create a Word document
    doc = Document()
    #doc.add_heading('PDF İçeriği', 0)
    doc.add_paragraph(pdf_content)
    
    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)   
    
    return send_file(doc_io, as_attachment=True, download_name='pdf_content.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

   
    
if __name__=='__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True, port=5000)
    




    